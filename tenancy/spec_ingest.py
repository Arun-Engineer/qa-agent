# tenancy/spec_ingest.py
from __future__ import annotations

import json
import os
import re
import uuid
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session
from sqlalchemy import select

from tenancy.content_models import SpecDocument, SpecChunk


ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".md"}
DEFAULT_MAX_UPLOAD_MB = int(os.getenv("MAX_UPLOAD_MB", "25"))

SPEC_FILES_DIR = Path(os.getenv("SPEC_FILES_DIR", "data/spec_files")).resolve()
SPEC_FILES_DIR.mkdir(parents=True, exist_ok=True)


def _uuid() -> str:
    return str(uuid.uuid4())


def _clean_text(t: str) -> str:
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # collapse extreme whitespace but keep paragraphing
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{4,}", "\n\n\n", t)
    return t.strip()


def chunk_text(raw_text: str, chunk_chars: int = 3500, overlap: int = 300) -> list[dict[str, Any]]:
    """
    Simple chunker: splits on paragraphs/headings then packs into ~chunk_chars with overlap.
    This is NOT token-perfect, but good enough to start.
    """
    text = _clean_text(raw_text)
    if not text:
        return []

    parts = re.split(r"\n{2,}", text)
    parts = [p.strip() for p in parts if p.strip()]

    chunks: list[str] = []
    buf = ""

    for p in parts:
        if len(buf) + len(p) + 2 <= chunk_chars:
            buf = (buf + "\n\n" + p).strip()
        else:
            if buf:
                chunks.append(buf)
            buf = p

    if buf:
        chunks.append(buf)

    # apply overlap by characters (rough)
    final: list[dict[str, Any]] = []
    cursor = 0
    for i, c in enumerate(chunks):
        start = max(0, cursor - overlap) if i > 0 else cursor
        # approximate end
        end = start + len(c)
        final.append(
            {
                "chunk_index": i,
                "content": c,
                "start_char": start,
                "end_char": end,
                "meta": {},
            }
        )
        cursor = end

    return final


def extract_text_from_bytes(filename: str | None, mime_type: str | None, data: bytes) -> str:
    name = (filename or "").lower()
    ext = Path(name).suffix.lower()

    # Plain text / md
    if ext in (".txt", ".md") or (mime_type or "").startswith("text/"):
        return data.decode("utf-8", errors="ignore")

    # DOCX
    if ext == ".docx":
        try:
            import docx  # python-docx
        except Exception as e:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e

        d = docx.Document(data)  # NOTE: python-docx expects a path or file-like. We'll fallback below.

    # PDF
    if ext == ".pdf" or (mime_type == "application/pdf"):
        try:
            import pdfplumber
        except Exception as e:
            raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber") from e

        # write temp then read (pdfplumber needs a file)
        tmp = SPEC_FILES_DIR / f"_tmp_{_uuid()}.pdf"
        tmp.write_bytes(data)
        try:
            out = []
            with pdfplumber.open(str(tmp)) as pdf:
                for page in pdf.pages:
                    out.append(page.extract_text() or "")
            return "\n\n".join(out)
        finally:
            try:
                tmp.unlink(missing_ok=True)
            except Exception:
                pass

    # DOCX fallback (python-docx needs file)
    if ext == ".docx":
        try:
            import docx
        except Exception as e:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e

        tmp = SPEC_FILES_DIR / f"_tmp_{_uuid()}.docx"
        tmp.write_bytes(data)
        try:
            d = docx.Document(str(tmp))
            out = []
            for p in d.paragraphs:
                if p.text:
                    out.append(p.text)
            return "\n".join(out)
        finally:
            try:
                tmp.unlink(missing_ok=True)
            except Exception:
                pass

    raise RuntimeError(f"Unsupported file type: {ext}. Allowed: {sorted(ALLOWED_EXTS)}")


def create_spec_from_text(
    db: Session,
    tenant_id: str,
    account_id: str | None,
    text: str,
    source: str = "paste",
    meta: dict[str, Any] | None = None,
) -> str:
    spec = SpecDocument(
        tenant_id=str(tenant_id),
        account_id=str(account_id) if account_id is not None else None,
        source=source,
        filename=None,
        mime_type="text/plain",
        raw_text=_clean_text(text),
        meta_json=json.dumps(meta or {}, ensure_ascii=False),
    )
    db.add(spec)
    db.commit()
    db.refresh(spec)

    chunks = chunk_text(spec.raw_text)
    for ch in chunks:
        db.add(
            SpecChunk(
                tenant_id=str(tenant_id),
                spec_id=spec.id,
                chunk_index=ch["chunk_index"],
                content=ch["content"],
                start_char=ch.get("start_char"),
                end_char=ch.get("end_char"),
                meta_json=json.dumps(ch.get("meta") or {}, ensure_ascii=False),
            )
        )
    db.commit()
    return spec.id


def create_spec_from_upload(
    db: Session,
    tenant_id: str,
    account_id: str | None,
    filename: str,
    mime_type: str | None,
    data: bytes,
    source: str = "upload",
    meta: dict[str, Any] | None = None,
    save_original: bool = True,
) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTS:
        raise ValueError(f"Unsupported file extension {ext}. Allowed: {sorted(ALLOWED_EXTS)}")

    if len(data) > DEFAULT_MAX_UPLOAD_MB * 1024 * 1024:
        raise ValueError(f"File too large. Max {DEFAULT_MAX_UPLOAD_MB}MB")

    raw_text = extract_text_from_bytes(filename, mime_type, data)
    raw_text = _clean_text(raw_text)

    spec = SpecDocument(
        tenant_id=str(tenant_id),
        account_id=str(account_id) if account_id is not None else None,
        source=source,
        filename=filename,
        mime_type=mime_type,
        raw_text=raw_text,
        meta_json=json.dumps(meta or {}, ensure_ascii=False),
    )
    db.add(spec)
    db.commit()
    db.refresh(spec)

    # save original file (optional)
    if save_original:
        out_dir = SPEC_FILES_DIR / str(tenant_id) / spec.id
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / filename).write_bytes(data)

    chunks = chunk_text(raw_text)
    for ch in chunks:
        db.add(
            SpecChunk(
                tenant_id=str(tenant_id),
                spec_id=spec.id,
                chunk_index=ch["chunk_index"],
                content=ch["content"],
                start_char=ch.get("start_char"),
                end_char=ch.get("end_char"),
                meta_json=json.dumps(ch.get("meta") or {}, ensure_ascii=False),
            )
        )
    db.commit()

    return spec.id


def get_spec(db: Session, tenant_id: str, spec_id: str) -> SpecDocument | None:
    return db.execute(
        select(SpecDocument).where(SpecDocument.tenant_id == str(tenant_id), SpecDocument.id == spec_id)
    ).scalar_one_or_none()


def get_chunks(db: Session, tenant_id: str, spec_id: str) -> list[SpecChunk]:
    rows = db.execute(
        select(SpecChunk)
        .where(SpecChunk.tenant_id == str(tenant_id), SpecChunk.spec_id == spec_id)
        .order_by(SpecChunk.chunk_index.asc())
    ).scalars().all()
    return list(rows)